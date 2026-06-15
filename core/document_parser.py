# =========================================
#  文档解析模块：TXT / PDF / DOCX -> Document 列表
#
#  功能：
#    - 支持解析 TXT、PDF、DOCX 格式文件
#    - 读取文件内容，转换为 Document 对象列表
#    - 异常捕获：文件不存在、格式不支持、文件损坏
#    - 适配 Windows 路径
#
#  使用方式：
#    from core.document_parser import parse_file, parse_directory
#    docs = parse_file("data/documents/test.pdf")
#    docs = parse_directory("data/documents/")
# =========================================
import os
import logging
from pathlib import Path
from typing import List, Optional, Set

# 日志
_logger = logging.getLogger(__name__)


# ============================================================
# Document 对象（简化版，兼容 LangChain 结构）
# ============================================================
class Document:
    """
    简化版文档对象，结构与 LangChain Document 一致。
    - page_content: 文本内容
    - metadata: 元数据（source、filename、page 等）
    """
    def __init__(self, page_content: str, metadata: Optional[dict] = None):
        self.page_content = page_content.strip() if page_content else ""
        self.metadata = metadata or {}

    def __repr__(self) -> str:
        source = self.metadata.get("source", "unknown")
        page = self.metadata.get("page", "")
        page_str = f", page={page}" if page else ""
        return f"Document(source={source}{page_str}, len={len(self.page_content)})"


# ============================================================
# 异常类
# ============================================================
class DocumentParseError(Exception):
    """文档解析异常统一封装。"""
    pass


class FileNotFoundError(DocumentParseError):
    """文件不存在。"""
    pass


class UnsupportedFormatError(DocumentParseError):
    """不支持的文件格式。"""
    pass


class FileCorruptedError(DocumentParseError):
    """文件损坏或无法读取。"""
    pass


# ============================================================
# 支持的文件扩展名（与 config/settings.yaml 中的 upload.allow_extensions 保持一致）
# ============================================================
DEFAULT_SUPPORTED_EXT: Set[str] = {".pdf", ".docx", ".doc", ".md", ".txt", ".markdown"}


# ============================================================
# 单文件解析
# ============================================================
def parse_file(
    file_path: str,
    supported_ext: Optional[Set[str]] = None,
) -> List[Document]:
    """
    解析单个文件，返回 Document 列表。

    参数：
        file_path: 文件路径（支持 Windows 绝对路径 / 相对路径）
        supported_ext: 支持的扩展名集合，默认使用 DEFAULT_SUPPORTED_EXT

    返回：
        List[Document]: 解析后的文档列表（PDF 每页一个 Document，其他格式整体一个）

    异常：
        FileNotFoundError: 文件不存在
        UnsupportedFormatError: 不支持的文件格式
        FileCorruptedError: 文件损坏或无法读取
    """
    # 1) 路径处理（适配 Windows）
    path = Path(file_path).resolve()  # 转为绝对路径
    ext = path.suffix.lower()

    # 2) 支持的扩展名
    ext_set = supported_ext or DEFAULT_SUPPORTED_EXT

    # 3) 文件存在性检查
    if not path.exists():
        raise FileNotFoundError(f"❌ 文件不存在：{path}")

    if not path.is_file():
        raise FileNotFoundError(f"❌ 路径不是文件：{path}")

    # 4) 格式支持检查
    if ext not in ext_set:
        raise UnsupportedFormatError(
            f"❌ 不支持的文件格式：{ext}\n"
            f"   当前支持：{sorted(ext_set)}\n"
            f"   文件路径：{path}"
        )

    # 5) 元数据基础字段
    meta = {
        "source": str(path),
        "filename": path.name,
        "extension": ext,
    }

    docs: List[Document] = []

    # --------------------------------------------------------
    # PDF 解析（PyPDF2）
    # --------------------------------------------------------
    if ext == ".pdf":
        try:
            from PyPDF2 import PdfReader
        except ImportError:
            raise DocumentParseError(
                "❌ 未安装 PyPDF2，请执行：pip install PyPDF2"
            )

        try:
            reader = PdfReader(str(path))
            # 检查是否加密
            if reader.is_encrypted:
                _logger.warning("PDF 文件已加密，尝试空密码解密：%s", path)
                try:
                    reader.decrypt("")  # 尝试空密码
                except Exception:
                    raise FileCorruptedError(
                        f"❌ PDF 文件已加密且无法解密：{path}\n"
                        f"   请提供未加密的 PDF 或使用正确密码。"
                    )

            # 逐页提取文本
            for idx, page in enumerate(reader.pages):
                try:
                    text = page.extract_text() or ""
                except Exception as e:
                    _logger.warning(
                        "PDF 页面 %d 提取失败（%s），跳过该页：%s",
                        idx + 1, str(e), path
                    )
                    continue

                if text.strip():
                    docs.append(Document(
                        text,
                        {**meta, "page": idx + 1, "total_pages": len(reader.pages)}
                    ))

            if not docs:
                _logger.warning("PDF 文件未提取到任何文本内容：%s", path)

        except FileCorruptedError:
            raise  # 直接抛出已封装的异常
        except Exception as e:
            raise FileCorruptedError(
                f"❌ PDF 文件损坏或无法读取：{path}\n"
                f"   错误详情：{type(e).__name__}: {e}"
            )

    # --------------------------------------------------------
    # DOCX 解析（python-docx）
    # --------------------------------------------------------
    elif ext in (".docx", ".doc"):
        # 注意：.doc（旧格式）python-docx 不支持，需要先转换为 .docx
        if ext == ".doc":
            _logger.warning(
                "⚠️  .doc（Word 97-2003）格式支持有限，建议转换为 .docx：%s",
                path
            )
            # 尝试按 .docx 读取，可能失败
            try:
                from docx import Document as DocxDocument
            except ImportError:
                raise DocumentParseError(
                    "❌ 未安装 python-docx，请执行：pip install python-docx"
                )
            try:
                doc = DocxDocument(str(path))
            except Exception as e:
                raise FileCorruptedError(
                    f"❌ .doc 文件无法读取（旧格式不支持）：{path}\n"
                    f"   请将文件另存为 .docx 格式后再上传。\n"
                    f"   错误详情：{type(e).__name__}: {e}"
                )
        else:
            try:
                from docx import Document as DocxDocument
            except ImportError:
                raise DocumentParseError(
                    "❌ 未安装 python-docx，请执行：pip install python-docx"
                )

        try:
            doc = DocxDocument(str(path))
            # 提取所有段落文本
            paragraphs = []
            for para in doc.paragraphs:
                text = para.text.strip()
                if text:
                    paragraphs.append(text)

            # 提取表格文本（可选）
            tables_text = []
            for table in doc.tables:
                for row in table.rows:
                    row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
                    if row_text:
                        tables_text.append(row_text)

            # 合并段落 + 表格
            all_text = []
            if paragraphs:
                all_text.extend(paragraphs)
            if tables_text:
                all_text.append("\n【表格内容】")
                all_text.extend(tables_text)

            full_text = "\n".join(all_text)
            if full_text.strip():
                docs.append(Document(full_text, meta))
            else:
                _logger.warning("DOCX 文件未提取到任何文本内容：%s", path)

        except Exception as e:
            raise FileCorruptedError(
                f"❌ DOCX 文件损坏或无法读取：{path}\n"
                f"   错误详情：{type(e).__name__}: {e}"
            )

    # --------------------------------------------------------
    # TXT / MD 解析（直接读取文本）
    # --------------------------------------------------------
    elif ext in (".txt", ".md", ".markdown"):
        try:
            # 尝试多种编码（UTF-8 优先，兼容 GBK）
            text = None
            for encoding in ("utf-8", "gbk", "gb2312", "latin-1"):
                try:
                    text = path.read_text(encoding=encoding)
                    break
                except UnicodeDecodeError:
                    continue

            if text is None:
                # 所有编码都失败，使用二进制 + errors="ignore"
                text = path.read_text(encoding="utf-8", errors="ignore")
                _logger.warning("TXT/MD 文件编码识别失败，使用忽略模式：%s", path)

            if text.strip():
                docs.append(Document(text, meta))
            else:
                _logger.warning("TXT/MD 文件内容为空：%s", path)

        except Exception as e:
            raise FileCorruptedError(
                f"❌ TXT/MD 文件无法读取：{path}\n"
                f"   错误详情：{type(e).__name__}: {e}"
            )

    # --------------------------------------------------------
    # 其他格式（理论上不会走到这里，因为前面已检查 ext）
    # --------------------------------------------------------
    else:
        raise UnsupportedFormatError(f"❌ 未实现的解析逻辑：{ext}")

    # 6) 返回结果
    _logger.info(
        "解析完成：%s -> %d 个 Document（总字符数：%d）",
        path.name, len(docs), sum(len(d.page_content) for d in docs)
    )
    return docs


# ============================================================
# 目录批量解析
# ============================================================
def parse_directory(
    dir_path: str,
    supported_ext: Optional[Set[str]] = None,
    recursive: bool = True,
) -> List[Document]:
    """
    递归解析目录下所有支持的文件。

    参数：
        dir_path: 目录路径
        supported_ext: 支持的扩展名集合
        recursive: 是否递归子目录（默认 True）

    返回：
        List[Document]: 所有文件的解析结果合并列表
    """
    path = Path(dir_path).resolve()

    if not path.exists():
        raise FileNotFoundError(f"❌ 目录不存在：{path}")

    if not path.is_dir():
        raise FileNotFoundError(f"❌ 路径不是目录：{path}")

    ext_set = supported_ext or DEFAULT_SUPPORTED_EXT
    all_docs: List[Document] = []
    file_count = 0
    error_count = 0

    # 遍历目录
    if recursive:
        walker = os.walk(str(path))
    else:
        # 不递归，只处理当前目录
        walker = [(str(path), [], [f for f in os.listdir(str(path)) if Path(f).is_file()])]

    for root, _, files in walker:
        for name in files:
            file_path = Path(root) / name
            ext = file_path.suffix.lower()

            if ext not in ext_set:
                continue  # 跳过不支持的格式

            try:
                docs = parse_file(str(file_path), supported_ext=ext_set)
                all_docs.extend(docs)
                file_count += 1
            except DocumentParseError as e:
                _logger.error("文件解析失败：%s\n%s", file_path, str(e))
                error_count += 1
                continue  # 单文件失败不影响整体流程

    _logger.info(
        "目录解析完成：%s -> %d 个文件成功，%d 个失败，共 %d 个 Document",
        path, file_count, error_count, len(all_docs)
    )

    if error_count > 0:
        _logger.warning(
            "⚠️  有 %d 个文件解析失败，请检查日志排查原因。",
            error_count
        )

    return all_docs


# ============================================================
# 工具函数：检查文件格式是否支持
# ============================================================
def is_supported(file_path: str, supported_ext: Optional[Set[str]] = None) -> bool:
    """检查文件格式是否在支持列表中。"""
    ext = Path(file_path).suffix.lower()
    ext_set = supported_ext or DEFAULT_SUPPORTED_EXT
    return ext in ext_set


def get_supported_extensions() -> Set[str]:
    """获取当前支持的文件扩展名集合。"""
    return DEFAULT_SUPPORTED_EXT.copy()


# ============================================================
# 导出
# ============================================================
__all__ = [
    "Document",
    "DocumentParseError",
    "FileNotFoundError",
    "UnsupportedFormatError",
    "FileCorruptedError",
    "parse_file",
    "parse_directory",
    "is_supported",
    "get_supported_extensions",
]


# ============================================================
# 直接运行本文件时的调试输出
# 用法：python -m core.document_parser
# ============================================================
if __name__ == "__main__":
    import sys

    logging.basicConfig(level=logging.INFO, format="%(levelname)s: %(message)s")

    # 测试：解析指定文件或目录
    if len(sys.argv) > 1:
        target = sys.argv[1]
        if Path(target).is_dir():
            docs = parse_directory(target)
        else:
            docs = parse_file(target)
        print(f"\n解析结果：{len(docs)} 个 Document")
        for i, d in enumerate(docs[:5]):  # 只打印前 5 个
            print(f"\n[{i+1}] {d}")
            print(d.page_content[:200] + "..." if len(d.page_content) > 200 else d.page_content)
    else:
        print("用法：python -m core.document_parser <文件或目录路径>")
        print(f"支持格式：{sorted(DEFAULT_SUPPORTED_EXT)}")